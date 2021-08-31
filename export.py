from docx import *
from docx.shared import Inches,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def sql_to_table(sql,name):

    #criação do arquivo docx 
    document = Document()

    #mudando as margens do documento
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    #imagem de cabeçalho centralizada (criada em um paragrafo que é depois centralizado com WD_PARAGRAPH ALIGMENT.CENTER)
    imagem = document.add_paragraph()
    imagem.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
    run = imagem.add_run("")
    run.add_picture("book_icon_v8_best_word.png",width= Inches(1.0))

    document.add_paragraph("Lista de Livros:")

    #criação da tabela
    table = document.add_table(1, 6)
    heading_cells = table.rows[0].cells
    heading_cells[0].text = ''
    heading_cells[1].text = 'Livro'
    heading_cells[2].text = 'Autor'
    heading_cells[3].text = 'Estilo'
    heading_cells[4].text = 'Dado a'
    heading_cells[5].text = 'Dedicado'
    table.style = "Light Grid"

    
    
    for i in range(len(sql)):
        celulas = table.add_row().cells
        for j in range(6):
            celulas[j].text = str(sql[i][j])
            
    for celulas in table.columns[0].cells:
        celulas.width = Cm(1.0)
    for celulas in table.columns[1].cells:
        celulas.width = Cm(9.0)
    for celulas in table.columns[2].cells:
        celulas.width = Cm(4.5)
    for celulas in table.columns[3].cells:
        celulas.width = Cm(3.0)
    for celulas in table.columns[4].cells:
        celulas.width = Cm(1.5)
    for celulas in table.columns[5].cells:
        celulas.width = Cm(1.5)
    
    # esses loops vão em cada célula e mudam a fonte 
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(8)
   
    document.save(str(name+".docx"))





